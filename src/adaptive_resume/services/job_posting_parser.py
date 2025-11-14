"""Job posting file parser service.

This module provides functionality to parse job postings from various file formats
(PDF, DOCX, TXT) and extract clean text for analysis.
"""

from typing import Optional, Tuple
import re
import logging
from pathlib import Path
from io import BytesIO

# PDF parsing
try:
    from pypdf import PdfReader
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

# DOCX parsing
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# Encoding detection
try:
    import chardet
    CHARDET_AVAILABLE = True
except ImportError:
    CHARDET_AVAILABLE = False

logger = logging.getLogger(__name__)


class JobPostingParserError(Exception):
    """Base exception for job posting parser errors."""
    pass


class UnsupportedFileTypeError(JobPostingParserError):
    """Raised when file type is not supported."""
    pass


class FileTooLargeError(JobPostingParserError):
    """Raised when file exceeds size limit."""
    pass


class FileParseError(JobPostingParserError):
    """Raised when file cannot be parsed."""
    pass


class JobPostingParser:
    """Service for parsing job postings from various file formats.

    Supports PDF, DOCX, and TXT files. Handles text extraction, cleaning,
    and normalization for downstream analysis.

    Attributes:
        max_file_size: Maximum file size in bytes (default: 10MB)
        supported_formats: Set of supported file extensions
    """

    # File size limit: 10MB
    MAX_FILE_SIZE = 10 * 1024 * 1024

    # Supported file extensions
    SUPPORTED_FORMATS = {'.pdf', '.docx', '.doc', '.txt'}

    # Common boilerplate patterns to remove
    BOILERPLATE_PATTERNS = [
        r'equal\s+opportunity\s+employer.*?(?:\.|$)',
        r'EOE.*?(?:\.|$)',
        r'we\s+are\s+an\s+equal\s+opportunity.*?(?:\.|$)',
        r'all\s+qualified\s+applicants.*?(?:\.|$)',
        r'reasonable\s+accommodations.*?(?:\.|$)',
        r'click\s+here\s+to\s+apply.*?(?:\.|$)',
        r'apply\s+now.*?(?:\.|$)',
    ]

    def __init__(self, max_file_size: Optional[int] = None):
        """Initialize parser with optional custom size limit.

        Args:
            max_file_size: Optional custom maximum file size in bytes
        """
        self.max_file_size = max_file_size or self.MAX_FILE_SIZE
        self.supported_formats = self.SUPPORTED_FORMATS.copy()

    def parse_file(self, file_path: str) -> str:
        """Parse a job posting file and extract text.

        Args:
            file_path: Path to the file to parse

        Returns:
            Cleaned text content from the file

        Raises:
            FileNotFoundError: If file doesn't exist
            UnsupportedFileTypeError: If file type is not supported
            FileTooLargeError: If file exceeds size limit
            FileParseError: If file cannot be parsed
        """
        path = Path(file_path)

        # Validate file exists
        if not path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")

        # Validate file size
        file_size = path.stat().st_size
        if file_size > self.max_file_size:
            raise FileTooLargeError(
                f"File size ({file_size} bytes) exceeds limit "
                f"({self.max_file_size} bytes)"
            )

        # Validate file type
        extension = path.suffix.lower()
        if extension not in self.supported_formats:
            raise UnsupportedFileTypeError(
                f"Unsupported file type: {extension}. "
                f"Supported types: {', '.join(self.supported_formats)}"
            )

        # Parse based on file type
        try:
            if extension == '.pdf':
                raw_text = self._parse_pdf(path)
            elif extension in {'.docx', '.doc'}:
                raw_text = self._parse_docx(path)
            elif extension == '.txt':
                raw_text = self._parse_txt(path)
            else:
                raise UnsupportedFileTypeError(f"Unsupported format: {extension}")

            # Clean and normalize text
            cleaned_text = self.clean_text(raw_text)

            return cleaned_text

        except (UnsupportedFileTypeError, FileTooLargeError):
            raise
        except Exception as e:
            logger.error(f"Error parsing file {file_path}: {e}")
            raise FileParseError(f"Failed to parse file: {str(e)}") from e

    def parse_text(self, text: str) -> str:
        """Parse and clean pasted text.

        Args:
            text: Raw text to clean

        Returns:
            Cleaned text content
        """
        return self.clean_text(text)

    def _parse_pdf(self, path: Path) -> str:
        """Extract text from PDF file.

        Args:
            path: Path to PDF file

        Returns:
            Extracted text

        Raises:
            FileParseError: If PDF parsing fails
        """
        if not PDF_AVAILABLE:
            raise FileParseError(
                "PDF support not available. Install with: pip install pypdf"
            )

        try:
            reader = PdfReader(str(path))
            text_parts = []

            for page in reader.pages:
                # Extract text from page
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)

            full_text = '\n\n'.join(text_parts)

            if not full_text.strip():
                raise FileParseError("PDF appears to be empty or image-based")

            return full_text

        except Exception as e:
            logger.error(f"PDF parsing error: {e}")
            raise FileParseError(f"Failed to parse PDF: {str(e)}") from e

    def _parse_docx(self, path: Path) -> str:
        """Extract text from DOCX file.

        Args:
            path: Path to DOCX file

        Returns:
            Extracted text

        Raises:
            FileParseError: If DOCX parsing fails
        """
        if not DOCX_AVAILABLE:
            raise FileParseError(
                "DOCX support not available. Install with: pip install python-docx"
            )

        try:
            doc = Document(str(path))
            text_parts = []

            # Extract paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)

            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = '\t'.join(cell.text.strip() for cell in row.cells)
                    if row_text.strip():
                        text_parts.append(row_text)

            full_text = '\n'.join(text_parts)

            if not full_text.strip():
                raise FileParseError("DOCX appears to be empty")

            return full_text

        except Exception as e:
            logger.error(f"DOCX parsing error: {e}")
            raise FileParseError(f"Failed to parse DOCX: {str(e)}") from e

    def _parse_txt(self, path: Path) -> str:
        """Extract text from TXT file with encoding detection.

        Args:
            path: Path to TXT file

        Returns:
            Extracted text

        Raises:
            FileParseError: If TXT reading fails
        """
        try:
            # Try to detect encoding
            encoding = self._detect_encoding(path)

            # Read file with detected encoding
            with open(path, 'r', encoding=encoding, errors='replace') as f:
                text = f.read()

            if not text.strip():
                raise FileParseError("TXT file appears to be empty")

            return text

        except Exception as e:
            logger.error(f"TXT parsing error: {e}")
            raise FileParseError(f"Failed to parse TXT: {str(e)}") from e

    def _detect_encoding(self, path: Path) -> str:
        """Detect file encoding using chardet.

        Args:
            path: Path to file

        Returns:
            Detected encoding (defaults to utf-8)
        """
        if not CHARDET_AVAILABLE:
            return 'utf-8'

        try:
            with open(path, 'rb') as f:
                raw_data = f.read()

            result = chardet.detect(raw_data)
            encoding = result.get('encoding', 'utf-8')
            confidence = result.get('confidence', 0)

            # Use utf-8 if confidence is too low
            if confidence < 0.7:
                logger.warning(
                    f"Low encoding confidence ({confidence}), using utf-8"
                )
                return 'utf-8'

            return encoding or 'utf-8'

        except Exception as e:
            logger.warning(f"Encoding detection failed: {e}, using utf-8")
            return 'utf-8'

    def clean_text(self, text: str) -> str:
        """Clean and normalize extracted text.

        Performs the following operations:
        - Normalizes whitespace
        - Removes boilerplate text (EOE statements, etc.)
        - Fixes common encoding issues
        - Removes excessive line breaks

        Args:
            text: Raw text to clean

        Returns:
            Cleaned text
        """
        if not text:
            return ""

        # Fix common encoding issues
        text = self._fix_encoding_issues(text)

        # Normalize line breaks
        text = text.replace('\r\n', '\n').replace('\r', '\n')

        # Remove excessive blank lines (more than 2 consecutive)
        text = re.sub(r'\n{3,}', '\n\n', text)

        # Normalize whitespace within lines
        lines = []
        for line in text.split('\n'):
            # Collapse multiple spaces to single space
            cleaned_line = re.sub(r'[ \t]+', ' ', line)
            lines.append(cleaned_line.strip())

        text = '\n'.join(lines)

        # Remove boilerplate text
        text = self._remove_boilerplate(text)

        # Final cleanup
        text = text.strip()

        return text

    def _fix_encoding_issues(self, text: str) -> str:
        """Fix common encoding issues in text.

        Args:
            text: Text with potential encoding issues

        Returns:
            Text with fixes applied
        """
        # Common replacements
        replacements = {
            '\u2019': "'",  # Right single quotation mark
            '\u2018': "'",  # Left single quotation mark
            '\u201c': '"',  # Left double quotation mark
            '\u201d': '"',  # Right double quotation mark
            '\u2013': '-',  # En dash
            '\u2014': '--', # Em dash
            '\u2026': '...', # Horizontal ellipsis
            '\xa0': ' ',    # Non-breaking space
            '\u00a0': ' ',  # Non-breaking space (alternate)
        }

        for old, new in replacements.items():
            text = text.replace(old, new)

        return text

    def _remove_boilerplate(self, text: str) -> str:
        """Remove common boilerplate text from job postings.

        Args:
            text: Text to clean

        Returns:
            Text with boilerplate removed
        """
        for pattern in self.BOILERPLATE_PATTERNS:
            text = re.sub(pattern, '', text, flags=re.IGNORECASE)

        return text

    def validate_file(self, file_path: str) -> Tuple[bool, Optional[str]]:
        """Validate a file without parsing it.

        Args:
            file_path: Path to file to validate

        Returns:
            Tuple of (is_valid, error_message)
            If valid, error_message is None
        """
        try:
            path = Path(file_path)

            if not path.exists():
                return False, "File not found"

            file_size = path.stat().st_size
            if file_size > self.max_file_size:
                size_mb = file_size / (1024 * 1024)
                limit_mb = self.max_file_size / (1024 * 1024)
                return False, f"File too large ({size_mb:.1f}MB, limit: {limit_mb:.1f}MB)"

            extension = path.suffix.lower()
            if extension not in self.supported_formats:
                return False, f"Unsupported file type: {extension}"

            return True, None

        except Exception as e:
            return False, str(e)

    @property
    def is_pdf_supported(self) -> bool:
        """Check if PDF parsing is available."""
        return PDF_AVAILABLE

    @property
    def is_docx_supported(self) -> bool:
        """Check if DOCX parsing is available."""
        return DOCX_AVAILABLE
